import pandas as pd
import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt

# ---------------------------
# INTRO SCREEN FOR FIRST TIME USERS
# ---------------------------
st.title("🚢 SPS–SOLAS Gap Analysis Tool")
st.markdown("""
Welcome to the **SPS–SOLAS Gap Analysis Tool**! 👋

This tool helps you evaluate the regulatory requirements when converting a vessel:
- From **cargo to SPS** with less or more than 60 special personnel
- From **SPS < 60 to SPS > 60**

### 📋 How to Use:
1. Use the **sidebar** to enter vessel info like GT, personnel, systems installed
2. Click **Run Gap Analysis**
3. View your compliance table and **download** Word/CSV reports

---
**Start by filling out the sidebar on the left ➡️**
""")

rules_info = {
    "SPS 2.2.3": ("Stability treated as cargo ship", "https://www.imo.org/en/OurWork/Safety/Pages/SpecialPurposeShips.aspx", "Review if special personnel > 60; may need to treat as passenger ship"),
    "SOLAS II-1/29.6.1.2": ("Auxiliary steering gear for ships ≤240 persons.", "https://www.imo.org/en/OurWork/Safety/Pages/SOLAS.aspx", "Review if vessel may exceed 240 personnel or lacks verified capacity"),
    "SOLAS II-1/29.6.1.1": ("Main steering gear for ships >240 persons.", "https://www.imo.org/en/OurWork/Safety/Pages/SOLAS.aspx", "Review if gear not confirmed for >240 personnel"),
    "SOLAS III": ("Life-saving appliances appropriate to personnel type.", "https://www.imo.org/en/OurWork/Safety/Pages/Life-Saving-Appliances.aspx", "Review for SPS vessels transitioning to >60 persons"),
    "SOLAS IV": ("GMDSS radio compliance for safety communication.", "https://www.imo.org/en/OurWork/Safety/Pages/Radio-Communications.aspx", ""),
    "SOLAS XI-2": ("Security plan for ship safety under ISPS Code.", "https://www.imo.org/en/OurWork/Security/Pages/SOLAS-XI-2.aspx", "Review if security measures in progress or partially implemented"),
    "SOLAS II-2": ("Fire protection systems compliance for vessel type.", "https://www.imo.org/en/OurWork/Safety/Pages/FireSafety.aspx", ""),
    "SOLAS II-1/19": ("Emergency electrical power supply standards.", "https://www.imo.org/en/OurWork/Safety/Pages/SOLAS.aspx", "Review if backup duration or independence not fully confirmed")
}

def gap_analysis(gt, sp, self_prop, ums, fire, lifeboat, emergency, steer, radio, security):
    if sp < 60:
        scenario = "Cargo to SPS <60"
    elif sp >= 60 and gt:
        scenario = "Cargo to SPS >60"
    else:
        scenario = "SPS <60 to SPS >60"

    results = []
    def check(rule, condition_type):
        desc, ref, note = rules_info[rule]
        if condition_type == "compliant":
            status = "✅ Compliant"
        elif condition_type == "review":
            status = "⚠️ Needs Review"
        else:
            status = "❌ Non-compliant"
        results.append({
            "Rule Regulation Number": rule,
            "Description of Rule": f"{desc} ({note})" if note else desc,
            "Regulatory Reference": ref,
            "Observation / Current Status": "",
            "Compliance or Not": status,
            "Audit Checklist Note": note if note else "Use rule description as guidance."
        })

    check("SPS 2.2.3", "compliant" if sp < 60 else "review")
    check("SOLAS II-1/29.6.1.2", "compliant" if steer == "II-1/29.6.1.2" else "review")
    check("SOLAS II-1/29.6.1.1", "compliant" if steer == "II-1/29.6.1.1" else "review")
    check("SOLAS III", "compliant" if lifeboat in ["cargo", "passenger"] else "non-compliant")
    check("SOLAS IV", "compliant" if radio else "non-compliant")
    check("SOLAS XI-2", "compliant" if security else "review")
    check("SOLAS II-2", "compliant" if fire else "non-compliant")
    check("SOLAS II-1/19", "compliant" if emergency else "review")

    df = pd.DataFrame(results)
    return scenario, df

def generate_summary(df):
    compliant = df['Compliance or Not'].str.contains("Compliant").sum()
    review = df['Compliance or Not'].str.contains("Review").sum()
    noncompliant = df['Compliance or Not'].str.contains("Non-compliant").sum()
    total = len(df)
    return f"Out of {total} rules checked: {compliant} are compliant, {review} need review, and {noncompliant} are non-compliant."

def export_to_word(scenario, df):
    doc = Document()
    doc.add_heading("SPS–SOLAS Gap Analysis Report", 0)
    doc.add_paragraph(f"Scenario: {scenario}")
    doc.add_paragraph(generate_summary(df))
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rule Regulation Number'
    hdr_cells[1].text = 'Description of Rule'
    hdr_cells[2].text = 'Compliance or Not'
    hdr_cells[3].text = 'Reference'

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = row['Rule Regulation Number']
        cells[1].text = row['Description of Rule']
        cells[2].text = row['Compliance or Not']
        cells[3].text = row['Regulatory Reference']

    filename = f"gap_analysis_{scenario.replace(' ', '_')}.docx"
    doc.save(filename)
    with open(filename, "rb") as file:
        st.download_button("📝 Download Word Report", file, file_name=filename)

# Sidebar inputs
with st.sidebar:
    gt = st.number_input("Gross Tonnage (GT)", min_value=0.0, value=500.0)
    sp = st.number_input("Special Personnel Count", min_value=0, value=50)
    self_prop = st.selectbox("Self Propelled?", ["Yes", "No"]) == "Yes"
    ums = st.selectbox("UMS Certified?", ["Yes", "No"]) == "Yes"
    fire = st.selectbox("Fire Protection Installed?", ["Yes", "No"]) == "Yes"
    lifeboat = st.selectbox("Lifeboat Type", ["cargo", "passenger", "none"])
    emergency = st.selectbox("Emergency Power Available?", ["Yes", "No"]) == "Yes"
    steer = st.selectbox("Steering Gear Compliance", ["II-1/29.6.1.1", "II-1/29.6.1.2", "none"])
    radio = st.selectbox("GMDSS Compliant?", ["Yes", "No"]) == "Yes"
    security = st.selectbox("Security Plan Onboard?", ["Yes", "No"]) == "Yes"

# Button to run analysis
if st.button("Run Gap Analysis"):
    scenario, df = gap_analysis(gt, sp, self_prop, ums, fire, lifeboat, emergency, steer, radio, security)
    st.subheader(f"Scenario: {scenario}")
    df_display = df.drop(columns=['Audit Checklist Note'])
    st.dataframe(df_display)
    st.markdown("**Summary:**")
    st.info(generate_summary(df))
    export_to_word(scenario, df)
